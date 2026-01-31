from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

#  REEMPLAZO INTELIGENTE DE MARCADORES
def reemplazar_marcador_en_parrafo(parrafo, marcador, valor):
    texto = "".join(run.text for run in parrafo.runs)
    if marcador not in texto:
        return False

    nuevo = texto.replace(marcador, valor)

    for run in parrafo.runs:
        run.text = ""

    parrafo.runs[0].text = nuevo
    return True


def reemplazar_texto(doc, marcador, valor):
    for p in doc.paragraphs:
        reemplazar_marcador_en_parrafo(p, marcador, valor)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    reemplazar_marcador_en_parrafo(p, marcador, valor)


#  INSERTAR TABLA DE PIEZAS
def insertar_tabla_piezas(doc, marcador, piezas):
    for p in doc.paragraphs:
        if marcador in p.text:
            p.text = ""

            # Crear tabla
            tabla = doc.add_table(rows=1, cols=2)

            # ===============================
            #  ALL BORDERS (todos los bordes)
            # ===============================
            tbl = tabla._element
            tblPr = tbl.get_or_add_tblPr()
            tblBorders = OxmlElement('w:tblBorders')

            for borde in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border_el = OxmlElement(f"w:{borde}")
                border_el.set(qn('w:val'), 'single')
                border_el.set(qn('w:sz'), '6')      # grosor
                border_el.set(qn('w:color'), '000000')  # negro
                tblBorders.append(border_el)

            tblPr.append(tblBorders)

            # ===============================
            #  ENCABEZADOS (gris claro + negritas + centrado)
            # ===============================
            hdr = tabla.rows[0].cells
            headers = ["Concepto", "Precio"]

            for i, texto in enumerate(headers):
                hdr[i].text = texto

                # Fondo gris claro
                sombreado = OxmlElement('w:shd')
                sombreado.set(qn('w:fill'), 'D9D9D9')  # gris claro
                hdr[i]._tc.get_or_add_tcPr().append(sombreado)

                # Negritas y centrado
                for paragraph in hdr[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True

            # ===============================
            #  FILAS DE PIEZAS
            # ===============================
            for nombre, precio in piezas:
                row = tabla.add_row().cells
                row[0].text = nombre
                row[1].text = f"$ {precio}"

                # Alinear verticalmente y centrar precio
                for c in row:
                    for paragraph in c.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Insertar tabla debajo del marcador
            p._element.addnext(tabla._element)
            break

# =====================================================
#   CONVERTIR TOTAL A LETRAS AUTOMÁTICAMENTE
# =====================================================
def convertir_total_a_letras(total):
    entero = int(total)
    centavos = int(round((total - entero) * 100))

    if centavos > 0:
        return f"{num2words(entero, lang='es')} pesos con {num2words(centavos, lang='es')} centavos"
    else:
        return f"{num2words(entero, lang='es')} pesos"


#  PROGRAMA PRINCIPAL
def main():

    print("=== GENERADOR DE PRESUPUESTOS ===")

    # IVA PRIMERO
    iva_porcentaje = float(input("IVA (%): "))

    # DATOS PRINCIPALES
    num_presupuesto = input("Número de presupuesto: ")
    fecha = input("Fecha (dd/mm/aaaa): ")
    dia, mes, anio = fecha.split("/")

    cliente = input("Cliente: ")
    vehiculo = input("Vehículo: ")
    np = input("N/P: ")
    serie = input("Número de serie: ")

    # NOTA DEL PRESUPUESTO
    nota = input("Nota del presupuesto: ")

    # PIEZAS
    piezas = []
    print("\n--- AGREGAR PIEZAS ---")
    print("Presiona ENTER para terminar.\n")

    while True:
        nombre = input("Pieza: ")
        if nombre.strip() == "":
            break
        precio = input(f"Precio de '{nombre}': ")
        piezas.append((nombre, precio))

    # CALCULAR TOTALES AUTOMÁTICAMENTE
    subtotal = sum(float(precio) for _, precio in piezas)
    iva = subtotal * (iva_porcentaje / 100)
    total = subtotal + iva

    # TEXTOS PARA EL DOCUMENTO
    subtotal_txt = f"{subtotal:.2f}"
    iva_txt = f"{iva:.2f}"
    total_txt = f"{total:.2f}"

    # *** TOTAL EN LETRAS AUTOMÁTICO ***
    total_letras = convertir_total_a_letras(total)

    # CARGAR PLANTILLA
    plantilla = "PlantillaPresupuesto.docx"
    if not os.path.exists(plantilla):
        print(f"No se encontró: {plantilla}")
        return

    doc = Document(plantilla)

    # REEMPLAZO DE MARCADORES
    reemplazar_texto(doc, "{NUM_PRESUPUESTO}", num_presupuesto)
    reemplazar_texto(doc, "{DIA}", dia)
    reemplazar_texto(doc, "{MES}", mes)
    reemplazar_texto(doc, "{ANIO}", anio)

    reemplazar_texto(doc, "{CLIENTE}", cliente)
    reemplazar_texto(doc, "{VEHICULO}", vehiculo)
    reemplazar_texto(doc, "{NP}", np)
    reemplazar_texto(doc, "{SERIE}", serie)

    reemplazar_texto(doc, "{NOTA}", nota)

    reemplazar_texto(doc, "{SUBTOTAL}", subtotal_txt)
    reemplazar_texto(doc, "{IVA}", iva_txt)
    reemplazar_texto(doc, "{TOTAL}", total_txt)
    reemplazar_texto(doc, "{TOTAL_LETRAS}", total_letras)

    # TABLA DE PIEZAS
    insertar_tabla_piezas(doc, "{TABLA_PIEZAS}", piezas)

    # GUARDAR ARCHIVO
    nombre_archivo = f"Presupuesto {num_presupuesto}.docx"
    doc.save(nombre_archivo)

    print(f"\nArchivo generado correctamente: {nombre_archivo}")


if __name__ == "__main__":
    main()
